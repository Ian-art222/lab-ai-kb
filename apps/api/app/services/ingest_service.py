from __future__ import annotations

import logging
import re
import hashlib
from datetime import datetime
from pathlib import Path

from PyPDF2 import PdfReader
from sqlalchemy import text
from sqlalchemy.orm import Session
from docx import Document

from app.core.config import settings as app_settings
from app.models.file_record import FileRecord
from app.models.knowledge import KnowledgeChunk
from app.models.system_setting import SystemSetting
from app.services import chunk_pipeline
from app.services.model_service import embed_texts
from app.services.settings_service import (
    build_embedding_index_standard,
    get_effective_embedding_batch_size,
)

CHUNK_SIZE = app_settings.ingest_chunk_size
CHUNK_OVERLAP = app_settings.ingest_chunk_overlap
MIN_CHUNK_CHARS = app_settings.ingest_min_chunk_chars
MIN_MEANINGFUL_TEXT_CHARS = 5
MAX_INDEX_TEXT_CHARS = app_settings.ingest_max_index_text_chars

logger = logging.getLogger(__name__)
PIPELINE_VERSION = "v2_phase2"
PIPELINE_VERSION_STRUCTURAL = chunk_pipeline.PIPELINE_VERSION


def _set_index_stage(db: Session, file_record: FileRecord, stage: str) -> None:
    file_record.index_status = stage
    file_record.pipeline_version = PIPELINE_VERSION
    db.commit()
    db.refresh(file_record)


def _resolve_file_path(file_record: FileRecord) -> Path:
    storage_path = file_record.storage_path or file_record.file_name
    return Path(app_settings.upload_dir) / storage_path


def _extract_text(file_record: FileRecord, file_path: Path) -> str:
    file_type = (file_record.file_type or "").lower()

    if file_type in {"txt", "md"}:
        return file_path.read_text(encoding="utf-8")
    if file_type in {"csv", "tsv"}:
        return file_path.read_text(encoding="utf-8")
    if file_type == "pdf":
        reader = PdfReader(str(file_path))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        return "\n\n".join([page for page in pages if page])
    if file_type == "docx":
        document = Document(str(file_path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        return "\n".join([paragraph for paragraph in paragraphs if paragraph])

    raise ValueError(f"暂不支持该文件类型索引：{file_type or 'unknown'}")


def _split_markdown_sections(text: str) -> list[dict]:
    lines = text.splitlines()
    sections: list[dict] = []
    current_title: str | None = None
    current_heading_level: int | None = None
    buf: list[str] = []
    in_code = False

    def flush() -> None:
        body = "\n".join(buf).strip()
        if not body:
            return
        block_type = "code" if body.startswith("```") else "paragraph"
        sections.append(
            {
                "text": body,
                "page_number": None,
                "section_title": current_title,
                "heading_level": current_heading_level,
                "block_type": block_type,
                "section_path": [current_title] if current_title else [],
                "parent_section": current_title,
                "source_type": "markdown",
            }
        )

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            buf.append(line)
            continue
        if not in_code and stripped.startswith("#"):
            flush()
            buf = []
            current_title = stripped.lstrip("#").strip() or current_title
            current_heading_level = len(stripped) - len(stripped.lstrip("#"))
            continue
        if not in_code and re.match(r"^\|.+\|$", stripped):
            buf.append(line)
            continue
        buf.append(line)
    flush()
    return sections


def _split_paragraph_segments(
    text: str,
    *,
    page_number: int | None,
    section_title: str | None = None,
) -> list[dict]:
    segments: list[dict] = []
    blocks = re.split(r"\n\s*\n", text.strip())
    for block in blocks:
        content = block.strip()
        if not content:
            continue
        block_type = "list" if re.match(r"^([-*]|\d+\.)\s+", content) else "paragraph"
        segments.append(
            {
                "text": content,
                "page_number": page_number,
                "section_title": section_title or _guess_section_title(content),
                "block_type": "table" if _looks_like_text_table(content) else block_type,
                "section_path": [section_title] if section_title else [],
                "parent_section": section_title,
                "source_type": "paragraph",
            }
        )
    return segments


def _looks_like_text_table(content: str) -> bool:
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
    if len(lines) < 2:
        return False
    pipe_lines = [ln for ln in lines if "|" in ln]
    if len(pipe_lines) >= 2:
        return True
    sep_like = [ln for ln in lines if re.match(r"^[\s\-+:|]{3,}$", ln)]
    return bool(sep_like and pipe_lines)


def _extract_segments(file_record: FileRecord, file_path: Path) -> list[dict]:
    file_type = (file_record.file_type or "").lower()

    if file_type in {"txt", "md"}:
        text = file_path.read_text(encoding="utf-8")
        if file_type == "md":
            md_segments = _split_markdown_sections(text)
            if md_segments:
                return md_segments
        return _split_paragraph_segments(text, page_number=None)
    if file_type in {"csv", "tsv"}:
        text = file_path.read_text(encoding="utf-8")
        rows = [line.strip() for line in text.splitlines() if line.strip()]
        return [
            {
                "text": row,
                "page_number": None,
                "section_title": "tabular_rows",
                "block_type": "table",
                "section_path": ["tabular_rows"],
                "parent_section": "tabular_rows",
                "source_type": file_type,
            }
            for row in rows
        ]
    if file_type == "pdf":
        reader = PdfReader(str(file_path))
        segments: list[dict] = []
        weak_pages = 0
        for index, page in enumerate(reader.pages):
            page_text = chunk_pipeline.normalize_pdf_page_text(page.extract_text() or "")
            if page_text:
                if len(page_text.strip()) < app_settings.ingest_pdf_min_chars_per_page:
                    weak_pages += 1
                segments.extend(
                    _split_paragraph_segments(
                        page_text,
                        page_number=index + 1,
                    )
                )
        if weak_pages and weak_pages >= max(1, len(reader.pages) // 2):
            logger.warning(
                "PDF text extraction looks weak for file_id=%s; likely scanned PDF without OCR support",
                file_record.id,
            )
        return segments
    if file_type == "docx":
        document = Document(str(file_path))
        segments: list[dict] = []
        current_heading: str | None = None
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (getattr(paragraph.style, "name", "") or "").lower()
            if style_name.startswith("heading"):
                current_heading = text
                continue
            segments.append(
                {
                    "text": text,
                    "page_number": None,
                    "section_title": current_heading or _guess_section_title(text),
                    "block_type": "paragraph",
                    "section_path": [current_heading] if current_heading else [],
                    "parent_section": current_heading,
                    "source_type": "docx",
                }
            )
        return segments

    raise ValueError(f"暂不支持该文件类型索引：{file_type or 'unknown'}")


def _guess_section_title(text: str) -> str | None:
    for line in text.splitlines():
        candidate = line.strip().strip("#").strip()
        if 4 <= len(candidate) <= 80:
            return candidate
    return None


def _limit_segments(segments: list[dict]) -> tuple[list[dict], bool]:
    limited: list[dict] = []
    total = 0
    truncated = False
    for segment in segments:
        content = segment["text"].strip()
        if not content:
            continue
        remaining = MAX_INDEX_TEXT_CHARS - total
        if remaining <= 0:
            truncated = True
            break
        if len(content) > remaining:
            limited.append(
                {
                    "text": content[:remaining],
                    "page_number": segment.get("page_number"),
                    "section_title": segment.get("section_title"),
                    "block_type": segment.get("block_type"),
                    "heading_level": segment.get("heading_level"),
                    "section_path": segment.get("section_path"),
                    "parent_section": segment.get("parent_section"),
                    "source_type": segment.get("source_type"),
                }
            )
            truncated = True
            break
        limited.append(
            {
                "text": content,
                "page_number": segment.get("page_number"),
                "section_title": segment.get("section_title"),
                "block_type": segment.get("block_type"),
                "heading_level": segment.get("heading_level"),
                "section_path": segment.get("section_path"),
                "parent_section": segment.get("parent_section"),
                "source_type": segment.get("source_type"),
            }
        )
        total += len(content)
    return limited, truncated


def _chunk_text(
    text: str,
    *,
    page_number: int | None = None,
    section_title: str | None = None,
    block_type: str | None = None,
) -> tuple[list[dict], str | None]:
    normalized = text.strip()
    if not normalized:
        return [], None
    if len(normalized) < MIN_MEANINGFUL_TEXT_CHARS:
        return [], "文本内容过短，无法建立稳定索引"

    chunks: list[dict] = []
    warning: str | None = None
    if len(normalized) < MIN_CHUNK_CHARS:
        return (
            [
                {
                    "content": normalized,
                    "page_number": page_number,
                    "section_title": section_title or _guess_section_title(normalized),
                    "block_type": block_type or "paragraph",
                }
            ],
            "文本较短，检索效果可能有限",
        )

    units = [u.strip() for u in re.split(r"\n\s*\n", normalized) if u.strip()]
    if not units:
        units = [normalized]

    buf = ""
    for unit in units:
        candidate = f"{buf}\n\n{unit}".strip() if buf else unit
        if len(candidate) <= CHUNK_SIZE:
            buf = candidate
            continue
        if buf:
            chunks.append(
                {
                    "content": buf,
                    "page_number": page_number,
                    "section_title": section_title or _guess_section_title(buf),
                    "block_type": block_type or "paragraph",
                }
            )
            overlap = buf[-CHUNK_OVERLAP:] if CHUNK_OVERLAP > 0 else ""
            buf = f"{overlap}\n{unit}".strip()
            if len(buf) > CHUNK_SIZE:
                buf = buf[:CHUNK_SIZE]
                warning = "存在超长段落，已按长度约束切分"
        else:
            chunks.append(
                {
                    "content": unit[:CHUNK_SIZE],
                    "page_number": page_number,
                    "section_title": section_title or _guess_section_title(unit),
                    "block_type": block_type or "paragraph",
                }
            )
            warning = "存在超长段落，已按长度约束切分"
            buf = unit[CHUNK_SIZE - CHUNK_OVERLAP : CHUNK_SIZE].strip()
    if buf and len(buf) >= MIN_CHUNK_CHARS:
        chunks.append(
            {
                "content": buf,
                "page_number": page_number,
                "section_title": section_title or _guess_section_title(buf),
                "block_type": block_type or "paragraph",
            }
        )

    return chunks, warning


def _mark_index_success(
    db: Session,
    *,
    file_record: FileRecord,
    extracted_text_length: int,
    warnings: list[str],
    indexed_at: datetime,
    index_embedding_provider: str | None,
    index_embedding_model: str | None,
    index_embedding_dimension: int | None,
    pipeline_version: str | None = None,
) -> FileRecord:
    file_record.extracted_text_length = extracted_text_length
    file_record.index_status = "partial_failed" if warnings else "indexed"
    file_record.indexed_at = indexed_at
    file_record.index_error = None
    file_record.last_error = None
    file_record.last_error_code = None
    file_record.index_warning = "；".join(warnings) if warnings else None
    file_record.index_embedding_provider = index_embedding_provider
    file_record.index_embedding_model = index_embedding_model
    file_record.index_embedding_dimension = index_embedding_dimension
    file_record.pipeline_version = pipeline_version or PIPELINE_VERSION
    db.commit()
    db.refresh(file_record)
    return file_record


def _mark_index_failure(
    db: Session,
    *,
    file_id: int,
    error_message: str,
) -> FileRecord:
    db.rollback()
    db.query(KnowledgeChunk).filter(KnowledgeChunk.file_id == file_id).delete()
    file_record = db.query(FileRecord).filter(FileRecord.id == file_id).first()
    if file_record is None:
        raise RuntimeError(error_message)
    file_record.index_status = "failed"
    file_record.indexed_at = None
    file_record.index_error = error_message
    file_record.last_error = error_message
    file_record.last_error_code = "internal_error"
    file_record.index_warning = None
    file_record.index_embedding_provider = None
    file_record.index_embedding_model = None
    file_record.index_embedding_dimension = None
    file_record.pipeline_version = PIPELINE_VERSION
    db.commit()
    db.refresh(file_record)
    return file_record


def ingest_file(db: Session, file_record: FileRecord) -> FileRecord:
    return ingest_file_job(db, file_record, prepare_indexing=True)


def ingest_file_job(
    db: Session,
    file_record: FileRecord,
    *,
    prepare_indexing: bool,
) -> FileRecord:
    if prepare_indexing:
        file_record.index_status = "pending"
        file_record.index_error = None
        file_record.index_warning = None
        file_record.indexed_at = None
        file_record.last_error = None
        file_record.last_error_code = None
        file_record.index_embedding_provider = None
        file_record.index_embedding_model = None
        file_record.index_embedding_dimension = None
        file_record.pipeline_version = (
            PIPELINE_VERSION_STRUCTURAL
            if app_settings.ingest_structural_chunking_enabled
            else PIPELINE_VERSION
        )
        db.commit()
        db.refresh(file_record)
    file_path = _resolve_file_path(file_record)

    try:
        if not file_path.exists():
            raise FileNotFoundError("物理文件不存在，无法建立索引")
        if file_record.file_size == 0:
            raise ValueError("文件为空，无法建立索引")

        text = _extract_text(file_record, file_path)
        _set_index_stage(db, file_record, "parsing")
        if not text.strip():
            raise ValueError("文件解析完成，但未提取到可用文本内容")

        _set_index_stage(db, file_record, "chunking")
        warnings: list[str] = []
        active_pipeline = PIPELINE_VERSION

        if app_settings.ingest_structural_chunking_enabled:
            rows_spec, pipe_warnings = chunk_pipeline.build_rows_spec(file_record, file_path)
            warnings.extend(pipe_warnings)
            active_pipeline = PIPELINE_VERSION_STRUCTURAL
        else:
            segments = _extract_segments(file_record, file_path)
            limited_segments, truncated = _limit_segments(segments)
            if truncated:
                warnings.append(f"文件文本较长，本次仅截取前 {MAX_INDEX_TEXT_CHARS} 个字符建立索引")

            rows_spec = []
            for seg_idx, segment in enumerate(limited_segments):
                seg_text = segment["text"].strip()
                if not seg_text:
                    continue
                page_no = segment.get("page_number")
                section_guess = segment.get("section_title") or _guess_section_title(seg_text)
                block_type = segment.get("block_type") or "paragraph"
                section_path = segment.get("section_path") or ([section_guess] if section_guess else [])
                source_type = segment.get("source_type") or (file_record.file_type or "").lower() or "text"
                parent_section = segment.get("parent_section") or section_guess
                seg_hash = hashlib.sha256(seg_text.encode("utf-8")).hexdigest()
                parent_meta = {
                    "doc_id": file_record.id,
                    "file_id": file_record.id,
                    "filename": file_record.file_name,
                    "source_file_name": file_record.file_name,
                    "page_number": page_no,
                    "page_range": [page_no, page_no] if page_no is not None else None,
                    "section_title": section_guess,
                    "section_path": section_path,
                    "parent_section": parent_section,
                    "segment_order": seg_idx,
                    "chunk_role": "parent",
                    "block_type": block_type,
                    "source_type": source_type,
                    "content_hash": seg_hash,
                    "pipeline_version": "v2_phase1",
                }
                parent_row_index = len(rows_spec)
                rows_spec.append(
                    {
                        "chunk_kind": "parent",
                        "parent_ref": None,
                        "content": seg_text,
                        "chunk_index": -1,
                        "page_number": page_no,
                        "section_title": section_guess,
                        "metadata_json": dict(parent_meta),
                    }
                )
                segment_chunks, segment_warning = _chunk_text(
                    seg_text,
                    page_number=page_no,
                    section_title=section_guess,
                    block_type=block_type,
                )
                for ci, ch in enumerate(segment_chunks or []):
                    child_meta = {
                        "doc_id": file_record.id,
                        "file_id": file_record.id,
                        "filename": file_record.file_name,
                        "source_file_name": file_record.file_name,
                        "page_number": ch.get("page_number"),
                        "page_range": [ch.get("page_number"), ch.get("page_number")]
                        if ch.get("page_number") is not None
                        else None,
                        "section_title": ch.get("section_title"),
                        "section_path": section_path,
                        "parent_section": parent_section,
                        "segment_order": seg_idx,
                        "child_index_in_segment": ci,
                        "chunk_role": "child",
                        "parent_row_index": parent_row_index,
                        "block_type": ch.get("block_type") or block_type,
                        "source_type": source_type,
                        "content_hash": hashlib.sha256(ch["content"].encode("utf-8")).hexdigest(),
                        "pipeline_version": "v2_phase1",
                    }
                    rows_spec.append(
                        {
                            "chunk_kind": "child",
                            "parent_ref": parent_row_index,
                            "content": ch["content"],
                            "chunk_index": -1,
                            "page_number": ch.get("page_number"),
                            "section_title": ch.get("section_title"),
                            "metadata_json": child_meta,
                        }
                    )
                if segment_warning and segment_warning not in warnings:
                    warnings.append(segment_warning)
        if not rows_spec:
            raise ValueError("文本内容不足以形成可用索引")

        for i, row in enumerate(rows_spec):
            row["chunk_index"] = i

        child_indices = [i for i, r in enumerate(rows_spec) if r["chunk_kind"] == "child"]

        logger.info(
            "Start ingest file_id=%s row_count=%s child_count=%s extracted_text_length=%s",
            file_record.id,
            len(rows_spec),
            len(child_indices),
            len(text),
        )

        db.query(KnowledgeChunk).filter(KnowledgeChunk.file_id == file_record.id).delete()

        settings = db.query(SystemSetting).filter(SystemSetting.id == 1).first()
        embeddings: list[list[float] | None] = []
        embedding_warning: str | None = None
        index_embedding_provider: str | None = None
        index_embedding_model: str | None = None
        index_embedding_dimension: int | None = None
        if (
            settings
            and settings.embedding_provider
            and settings.embedding_api_base
            and settings.embedding_api_key
            and settings.embedding_model
        ):
            current_index_standard = build_embedding_index_standard(
                embedding_provider=settings.embedding_provider,
                embedding_model=settings.embedding_model,
            )
            effective_batch = get_effective_embedding_batch_size(
                embedding_provider_raw=settings.embedding_provider,
                db_batch_size=settings.embedding_batch_size,
            )
            total_batches = (
                (len(child_indices) + effective_batch - 1) // effective_batch if child_indices else 0
            )
            logger.info(
                "Embedding ingest file_id=%s provider=%s model=%s index_standard=%s child_chunk_count=%s total_batches=%s "
                "effective_batch_size=%s db_embedding_batch_size=%s",
                file_record.id,
                settings.embedding_provider,
                settings.embedding_model,
                current_index_standard,
                len(child_indices),
                total_batches,
                effective_batch,
                settings.embedding_batch_size,
            )
            if child_indices:
                _set_index_stage(db, file_record, "embedding")
                embeddings = embed_texts(
                    provider=settings.embedding_provider,
                    api_base=settings.embedding_api_base,
                    api_key=settings.embedding_api_key,
                    model=settings.embedding_model,
                    inputs=[rows_spec[i]["content"] for i in child_indices],
                    embedding_batch_size_from_db=settings.embedding_batch_size,
                )
            if embeddings and embeddings[0]:
                index_embedding_provider = settings.embedding_provider
                index_embedding_model = settings.embedding_model
                index_embedding_dimension = len(embeddings[0])
            logger.info(
                "Embedding ingest completed file_id=%s embedding_count=%s index_standard=%s embedding_dim=%s",
                file_record.id,
                len(embeddings),
                current_index_standard,
                index_embedding_dimension,
            )
        else:
            embedding_warning = "Embedding 配置不完整，本次仅完成文本分块索引"
            warnings.append(embedding_warning)

        emb_by_row: dict[int, list[float] | None] = {
            row_i: embeddings[j] if j < len(embeddings) else None
            for j, row_i in enumerate(child_indices)
        }

        now = datetime.utcnow()
        # parent_ref = rows_spec index of the parent row → DB id after phase 1.
        parent_db_id_by_spec_index: dict[int, int] = {}

        # Phase 1: parents only (no embedding); flush so children get stable parent_chunk_id.
        parent_rows: list[tuple[int, KnowledgeChunk]] = []
        for idx, spec in enumerate(rows_spec):
            if spec["chunk_kind"] != "parent":
                continue
            meta = spec.get("metadata_json") or {}
            row = KnowledgeChunk(
                file_id=file_record.id,
                folder_id=file_record.folder_id,
                chunk_index=spec["chunk_index"],
                content=spec["content"],
                section_title=spec.get("section_title"),
                page_number=spec.get("page_number"),
                token_count=meta.get("approx_tokens") or chunk_pipeline.approx_tokens(spec["content"]),
                embedding=None,
                embedding_vec=None,
                chunk_kind=spec["chunk_kind"],
                parent_chunk_id=None,
                metadata_json=spec.get("metadata_json"),
                created_at=now,
                updated_at=now,
            )
            db.add(row)
            parent_rows.append((idx, row))
        db.flush()
        for idx, row in parent_rows:
            parent_db_id_by_spec_index[idx] = row.id

        # Phase 2: children with embeddings and parent_chunk_id from parent_ref (rows_spec index).
        for spec in rows_spec:
            if spec["chunk_kind"] != "child":
                continue
            pref = spec["parent_ref"]
            parent_db_id = (
                parent_db_id_by_spec_index.get(pref) if pref is not None else None
            )
            emb = emb_by_row.get(spec["chunk_index"])
            ev = None
            if emb and len(emb) == app_settings.qa_pgvector_dimensions:
                ev = emb
            cmeta = spec.get("metadata_json") or {}
            row = KnowledgeChunk(
                file_id=file_record.id,
                folder_id=file_record.folder_id,
                chunk_index=spec["chunk_index"],
                content=spec["content"],
                section_title=spec.get("section_title"),
                page_number=spec.get("page_number"),
                token_count=cmeta.get("approx_tokens") or chunk_pipeline.approx_tokens(spec["content"]),
                embedding=emb,
                embedding_vec=ev,
                chunk_kind=spec["chunk_kind"],
                parent_chunk_id=parent_db_id,
                metadata_json=spec.get("metadata_json"),
                created_at=now,
                updated_at=now,
            )
            db.add(row)
        db.flush()

        # Lexical index only on retrieval rows (child / legacy); parents skip FTS.
        db.execute(
            text(
                "UPDATE knowledge_chunks "
                "SET search_vector = to_tsvector('simple', coalesce(content, '')) "
                "WHERE file_id = :fid AND (chunk_kind = 'child' OR chunk_kind IS NULL)"
            ),
            {"fid": file_record.id},
        )

        return _mark_index_success(
            db,
            file_record=file_record,
            extracted_text_length=len(text),
            warnings=warnings,
            indexed_at=now,
            index_embedding_provider=index_embedding_provider,
            index_embedding_model=index_embedding_model,
            index_embedding_dimension=index_embedding_dimension,
            pipeline_version=active_pipeline,
        )
    except Exception as exc:
        logger.exception(
            "Ingest failed file_id=%s content_hash=%s error=%s",
            file_record.id,
            file_record.content_hash,
            exc,
        )
        return _mark_index_failure(
            db,
            file_id=file_record.id,
            error_message=str(exc),
        )
